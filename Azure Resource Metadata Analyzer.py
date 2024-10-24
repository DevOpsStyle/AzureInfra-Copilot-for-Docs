import os
import docx
import csv
import requests
import collections.abc
from azure.identity import ClientSecretCredential
from azure.mgmt.resource import ResourceManagementClient

# Informazioni del Service Principal (che hai ottenuto dal comando az ad sp create-for-rbac)
tenant_id = '<replace with tenant id'
client_id = '<replace with app client id>'
client_secret = '<replace with secret>'

# Configurazione degli endpoint e delle chiavi API di OpenAI
API_KEY = "<replace with api key for OpenAI>"
QUESTION_ENDPOINT = "<replace with OpenAI endpoint"

# ID della sottoscrizione Azure
subscription_id = '<replace with the landing zone subid where the workload is placed>'

# Autenticazione tramite il Service Principal
credential = ClientSecretCredential(
    tenant_id=tenant_id,
    client_id=client_id,
    client_secret=client_secret
)

# Crea il client per la gestione delle risorse di Azure
resource_client = ResourceManagementClient(credential, subscription_id)

# Funzione per appiattire un dizionario annidato in una struttura piana
def flatten_dict(d, parent_key='', sep='_'):
    items = []
    for k, v in d.items():
        new_key = f"{parent_key}{sep}{k}" if parent_key else k

        # Se il valore è un dizionario, lo appiattiamo ricorsivamente
        if isinstance(v, collections.abc.MutableMapping):
            items.extend(flatten_dict(v, new_key, sep=sep).items())
        # Se il valore è una lista, appiattiamo ogni elemento della lista con un suffisso numerico
        elif isinstance(v, list):
            for i, item in enumerate(v):
                # Controlliamo se l'elemento nella lista è a sua volta un dizionario
                if isinstance(item, collections.abc.MutableMapping):
                    items.extend(flatten_dict(item, f"{new_key}{sep}{i}", sep=sep).items())
                else:
                    items.append((f"{new_key}{sep}{i}", item))
        # Se il valore non è un dizionario né una lista, lo aggiungiamo direttamente
        else:
            items.append((new_key, v))
    return dict(items)


# Funzione per ottenere tutte le risorse con un determinato tag
def get_resources_by_tag(tag_key, tag_value):
    tag_filter = f"tagName eq '{tag_key}' and tagValue eq '{tag_value}'"
    print(f"Cercando risorse con il tag: {tag_key}={tag_value}")

    resources = resource_client.resources.list(filter=tag_filter)
    resources_list = []

    for resource in resources:
        resource_details = {
            'name': resource.name,
            'id': resource.id,
            'location': resource.location,
            'type': resource.type,
            'tags': resource.tags
        }
        resources_list.append(resource_details)
        print(f"Trovata risorsa: {resource.name} - {resource.type}")

    return resources_list


# Funzione per ottenere i Resource Groups con un determinato tag
def get_resource_groups_by_tag(tag_key, tag_value):
    print(f"Cercando resource groups con il tag: {tag_key}={tag_value}")
    resource_groups = resource_client.resource_groups.list()
    matching_resource_groups = []

    for rg in resource_groups:
        if rg.tags and tag_key in rg.tags and rg.tags[tag_key] == tag_value:
            print(f"Trovato resource group: {rg.name}")
            matching_resource_groups.append(rg)

    return matching_resource_groups


# Funzione per ottenere le risorse all'interno di un Resource Group
def get_resources_in_resource_group(resource_group_name):
    print(f"Recuperando risorse dal resource group: {resource_group_name}")
    resources = resource_client.resources.list_by_resource_group(resource_group_name)
    resources_list = []

    for resource in resources:
        resource_details = {
            'name': resource.name,
            'id': resource.id,
            'location': resource.location,
            'type': resource.type,
            'tags': resource.tags
        }
        resources_list.append(resource_details)
        print(f"Trovata risorsa nel resource group {resource_group_name}: {resource.name} - {resource.type}")

    return resources_list


# Funzione per ottenere l'API più recente per una risorsa specifica
def get_latest_api_version(resource_type):
    provider_namespace, resource_type_name = resource_type.split('/', 1)
    provider = resource_client.providers.get(provider_namespace)
    resource_type_info = next(
        (t for t in provider.resource_types if t.resource_type == resource_type_name), None
    )
    if resource_type_info:
        return sorted(resource_type_info.api_versions, reverse=True)[0]
    return None


# Funzione per ottenere i metadati completi di una risorsa specifica
def get_resource_metadata(resource):
    api_version = get_latest_api_version(resource['type'])
    if api_version:
        print(f"Usando l'API version: {api_version} per la risorsa {resource['name']}")
        resource_metadata = resource_client.resources.get_by_id(resource['id'], api_version=api_version)
        return resource_metadata
    else:
        print(f"Impossibile trovare l'API per la risorsa {resource['name']} con tipo {resource['type']}")
        return None


# Funzione per generare la overview del workload leggendo il file CSV
def generate_workload_overview():
    resources_info = []

    with open("resources_with_expanded_metadata.csv", 'r', encoding='utf-8') as csv_file:
        reader = csv.DictReader(csv_file)
        for row in reader:
            resources_info.append(row)

    resources_str = "\n".join([
        f"Name: {row['Name']}, Type: {row['Type']}, Location: {row['Location']}, Resource Group: {row['Resource ID'].split('/')[4]}"
        for row in resources_info])

    payload = {
        "messages": [
            {"role": "system",
             "content": "You are an expert Azure Architect and Documentation Writer. Your job is to create a clear and detailed overview of an Azure workload."},
            {"role": "user",
             "content": f"Here is the list of resources in the workload:\n{resources_str}.\nGenerate a detailed and human-readable overview."}
        ],
        "temperature": 0.7,
        "max_tokens": 16000
    }

    response = requests.post(
        QUESTION_ENDPOINT,
        headers={
            "Content-Type": "application/json",
            "api-key": API_KEY
        },
        json=payload
    )

    response.raise_for_status()
    response_from_copilot = response.json()['choices'][0]['message']['content'].strip()

    return response_from_copilot


# Funzione per generare la documentazione con OpenAI
def generate_infra_config(metadata_list):
    document_content = ""

    for metadata in metadata_list:
        metadata_str = str(metadata)

        payload = {
            "messages": [
                {"role": "system",
                 "content": "You are an expert Azure Architect and Documentation Writer."},
                {"role": "user",
                 "content": f"Here is the metadata for an Azure resource: \n{metadata_str}.\nPlease generate a detailed and human-readable documentation."}
            ],
            "temperature": 0.7,
            "max_tokens": 16000
        }

        response = requests.post(
            QUESTION_ENDPOINT,
            headers={
                "Content-Type": "application/json",
                "api-key": API_KEY
            },
            json=payload
        )

        response.raise_for_status()
        response_from_copilot = response.json()['choices'][0]['message']['content'].strip()

        document_content += response_from_copilot + "\n\n"

    with open("architecture.txt", "a") as architecturefile:
        architecturefile.write(document_content)

    return document_content


# Funzione per convertire il file txt in docx e aggiungere l'overview
def txt_to_docx():
    print("Generazione dei file in corso...")

    doc = docx.Document()

    # Genera l'overview del workload
    overview = generate_workload_overview()

    # Aggiungi "Workload Overview" come Titolo 1
    doc.add_heading("Workload Overview", level=1)
    doc.add_paragraph(overview)

    # Aggiungi un'interruzione di pagina per iniziare i dettagli su una nuova pagina
    doc.add_page_break()

    # Aggiungi "Workload Details" come Titolo 1
    doc.add_heading("Workload Details", level=1)

    # Aggiungi il contenuto del file architecture.txt
    with open("architecture.txt", 'r', encoding='utf-8', errors='ignore') as openfile:
        line = openfile.read()
        doc.add_paragraph(line)

    # Salva il documento Word
    doc.save("Output.docx")
    print("Il file Output.docx è stato creato con successo.")


# Funzione per eliminare il file architecture.txt al termine del run
def cleanup_files():
    if os.path.exists("architecture.txt"):
        os.remove("architecture.txt")
        print("File architecture.txt eliminato.")
    else:
        print("File architecture.txt non trovato, nessuna eliminazione necessaria.")


# Funzione per salvare le risorse con i metadati dinamici in un file CSV
def save_resources_with_expanded_metadata_to_csv(resources, metadata_list):
    all_keys = set()

    # Raccogli tutte le chiavi disponibili nei metadati, incluse quelle dei dizionari annidati
    for metadata in metadata_list:
        flat_metadata = flatten_dict(metadata.__dict__)  # Appiattiamo il dizionario dei metadati
        all_keys.update(flat_metadata.keys())

    # Trasforma il set delle chiavi in una lista ordinata per mantenere ordine nelle colonne
    all_keys = list(all_keys)

    # Scrivi le risorse e i loro metadati in un file CSV
    with open("resources_with_expanded_metadata.csv", mode="w", newline='', encoding="utf-8") as csv_file:
        # Creiamo l'header del CSV con tutte le chiavi uniche
        fieldnames = ['Name', 'Resource ID', 'Location', 'Type', 'Tags'] + all_keys
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames)

        # Scriviamo l'intestazione
        writer.writeheader()

        # Scrivi i dettagli di ogni risorsa e i suoi metadati
        for resource, metadata in zip(resources, metadata_list):
            resource_row = {
                'Name': resource['name'],
                'Resource ID': resource['id'],
                'Location': resource['location'],
                'Type': resource['type'],
                'Tags': resource['tags']
            }

            # Appiattiamo i metadati annidati prima di scriverli nel CSV
            flat_metadata = flatten_dict(metadata.__dict__)

            # Aggiungi i metadati alle colonne, gestendo eventuali valori mancanti
            for key in all_keys:
                resource_row[key] = flat_metadata.get(key, 'N/A')  # Usa 'N/A' per valori mancanti

            # Scrivi la riga nel CSV
            writer.writerow(resource_row)

    print("Il file resources_with_expanded_metadata.csv è stato creato con successo.")


# Funzione principale per orchestrare tutte le operazioni
def main():
    if os.path.exists("Output.docx") or os.path.exists("resources_with_expanded_metadata.csv"):
        print("I file Output.docx o resources_with_expanded_metadata.csv esistono già. Lo script non verrà eseguito.")
        return

    tag_key = 'Workload'
    tag_value = 'Production'

    all_resources = []
    added_resource_ids = set()

    resources = get_resources_by_tag(tag_key, tag_value)

    for resource in resources:
        if resource['id'] not in added_resource_ids:
            all_resources.append(resource)
            added_resource_ids.add(resource['id'])

    resource_groups = get_resource_groups_by_tag(tag_key, tag_value)

    for rg in resource_groups:
        rg_resources = get_resources_in_resource_group(rg.name)

        for resource in rg_resources:
            if resource['id'] not in added_resource_ids:
                all_resources.append(resource)
                added_resource_ids.add(resource['id'])

    # Ottieni i metadati di tutte le risorse
    metadata_list = []
    for resource in all_resources:
        print(f"\nRecuperando i metadati per la risorsa: {resource['name']}")
        metadata = get_resource_metadata(resource)
        if metadata:
            metadata_list.append(metadata)

    # Salva tutte le risorse (dirette e dai resource group) con i metadati in un file CSV
    save_resources_with_expanded_metadata_to_csv(all_resources, metadata_list)

    # Genera la documentazione utilizzando i metadati e OpenAI
    generate_infra_config(metadata_list)

    # Converte il file txt generato in un documento Word con la overview
    txt_to_docx()

    # Elimina il file architecture.txt al termine del run
    cleanup_files()


if __name__ == '__main__':
    main()
